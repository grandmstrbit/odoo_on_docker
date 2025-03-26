from datetime import date, datetime
from dateutil.relativedelta import relativedelta
import base64
import io

from odoo import api, fields, models
from odoo.exceptions import UserError, ValidationError
from odoo.tools.float_utils import float_compare, float_is_zero
from odoo.http import request, content_disposition
from werkzeug.wrappers import Response

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.oxml import OxmlElement, ns


class SroContactsWork(models.Model):
    _name = 'sro.contacts.work'
    _description = "Наличие права на выполнение работ"

    has_work_rights = fields.Boolean(string="Правом не наделен")
    
    number = fields.Char(string="№")
    right_status = fields.Selection([
        ('active', 'Действует'),
        ('suspended', 'Прекращено'),
    ], string="Статус права")
    right_effective_date = fields.Date(string="Дата вступления решения в силу")
    right_basis = fields.Many2one('blog.post', string="Основание выдачи (Документ №)") 
    right_basis_url = fields.Html(string="Основание выдачи (Документ №)", compute="_compute_right_basis_url", store=True)
    construction_object = fields.Selection([
        ('yes', 'Да'),
        ('no', 'Нет'),
    ], string="Объект капитального строительства")
    hazardous_objects = fields.Selection([
        ('yes', 'Да'),
        ('no', 'Нет'),
    ], string="Особо опасные, технически сложные и уникальные объекты")
    nuclear_objects = fields.Selection([
        ('yes', 'Да'),
        ('no', 'Нет'),
    ], string="Объекты использования атомной энергии")
    odo_right = fields.Selection([
        ('yes', 'Да'),
        ('no', 'Нет'),
        ('draft', 'Правом не наделен'),
    ], string="Наличие права на ОДО")
    hide_odo_date = fields.Date()
    hide_odo_doc = fields.Many2one('blog.post')
    hide_odo_doc_link = fields.Html(compute="_compute_hide_doc_link", store=True) 
    show_hide_fields = fields.Boolean(compute="_compute_show_hide_fields", store=True)
    odo_combined_info = fields.Html(compute="_compute_odo_combined_info", string="Наличие права на ОДО")
    hazardous_date = fields.Date()
    combined_hazardous = fields.Html(compute="_compute_combined_hazardous", string="Особо опасные, технически сложные и уникальные объекты")

    partner2_id = fields.Many2one('res.partner', string="Выполнение работ", invisible=True)
    
    @api.onchange('has_work_rights')
    def _onchange_has_work_rights(self):
        """Очищает поля, если галочка установлена"""
        if self.has_work_rights:
            self.update({
                'number': False,
                'right_status': False,
                'right_effective_date': False,
                'right_basis': False,
                'construction_object': False,
                'hazardous_objects': False,
                'nuclear_objects': False,
                'odo_right': 'draft',
            })

    @api.depends("right_basis")
    def _compute_right_basis_url(self):
        for rec in self:
            if rec.right_basis:
                url = f"/blog/resheniia-pravleniia-3/{rec.right_basis.id}"
                rec.right_basis_url = f'<a href="{url}" target="_blank">{rec.right_basis.name}</a>'
            else:
                rec.right_basis_url = ""

    @api.depends("hide_odo_doc")
    def _compute_hide_doc_link(self):
        for rec in self:
            if rec.hide_odo_doc:
                url = f"/blog/resheniia-pravleniia-3/{rec.hide_odo_doc.id}"
                rec.hide_odo_doc_link = f'<a href="{url}" target="_blank">{rec.hide_odo_doc.name}</a>'
            else:
                rec.hide_odo_doc_link = ""

    @api.depends("odo_right")
    def _compute_show_hide_fields(self):
        for rec in self:
            rec.show_hide_fields = rec.odo_right in ['yes', 'no']

    @api.depends("odo_right", "hide_odo_date", "hide_odo_doc_link")
    def _compute_odo_combined_info(self):
        for rec in self:
            if rec.odo_right in ['yes', 'no', 'draft']:
                if rec.odo_right == 'yes':
                    color_class = 'text-success'
                elif rec.odo_right == 'no':
                    color_class = 'text-danger'
                else:
                    color_class = ''
                date_str = rec.hide_odo_date.strftime("%d.%m.%Y") if rec.hide_odo_date else ""
                link = rec.hide_odo_doc_link or ""

                rec.odo_combined_info = f'<span class="{color_class}">{dict(self._fields["odo_right"].selection).get(rec.odo_right, "")}</span><br>{date_str}<br>{link}'
            else:
                rec.odo_combined_info = ""

    @api.depends("hazardous_objects", "hazardous_date")
    def _compute_combined_hazardous(self):
        for rec in self:
            if rec.hazardous_objects in ['yes', 'no']:
                color_class = 'text-success' if rec.hazardous_objects == 'yes' else 'text-danger'
                date_str = rec.hazardous_date.strftime("%d.%m.%Y") if rec.hazardous_date else ""

                rec.combined_hazardous = f'<span class="{color_class}">{dict(self._fields["hazardous_objects"].selection).get(rec.hazardous_objects, "")}</span><br>{date_str}'
            else:
                rec.combined_hazardous = ""

    def action_export_work_docx(self):
        partner = self.mapped('partner2_id')

        if not partner:
            raise UserError("Не удалось определить партнера!")
        records = self.search([('partner2_id', '=', partner.id)])

        if not records:
            raise UserError("Нет данных для экспорта!")

        doc = Document()
        section = doc.sections[0]

        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        column_widths = [Cm(1), Cm(3), Cm(5), Cm(4), Cm(4), Cm(4), Cm(4), Cm(4)]

        table = doc.add_table(rows=1, cols=len(column_widths))

        for row in table.rows:
            for cell in row.cells:
                for border in ['top', 'left', 'bottom', 'right']:
                    cell._element.get_or_add_tcPr().append(OxmlElement(f'w:{border}'))

        for i, cell in enumerate(table.rows[0].cells):
            cell.width = column_widths[i]

        def set_font(cell, bold=False):
            for para in cell.paragraphs:
                if not para.runs:
                    para.add_run()
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = bold

        headers = ["№", "Статус права", "Дата вступления решения в силу", "Основание выдачи (Документ №)", "Объект капитального строительства", "Особо опасные, технически сложные и уникальные объекты", "Объекты использования атомной энергии", "Наличие права на ОДО (Дата вступления решения в силу, Документ №)"]
        
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            set_font(cell, bold=True)

        for record in records:
            row_cells = table.add_row().cells
            row_cells[0].text = record.number or "—"
            row_cells[1].text = dict(record._fields['right_status'].selection).get(record.right_status, "—")
            row_cells[2].text = record.right_effective_date.strftime('%d.%m.%Y') if record.right_effective_date else "—"
            row_cells[3].text = record.right_basis.name or "—"
            row_cells[4].text = dict(record._fields['construction_object'].selection).get(record.construction_object, "—")
            row_cells[5].text = dict(record._fields['hazardous_objects'].selection).get(record.hazardous_objects, "—")
            row_cells[6].text = dict(record._fields['nuclear_objects'].selection).get(record.nuclear_objects, "—")

            odo_text = dict(record._fields['odo_right'].selection).get(record.odo_right, "—")
            odo_date = record.hide_odo_date.strftime('%d.%m.%Y') if record.hide_odo_date else "—"
            odo_link = record.hide_odo_doc.name if record.hide_odo_doc else "—"  # Имя документа вместо HTML-ссылки

            row_cells[7].text = f"{odo_text},\n{odo_date},\n{odo_link}"

            for i, cell in enumerate(row_cells):
                cell.width = column_widths[i]
                set_font(cell)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        registration_number = self.partner2_id.registration_number or "записи"
        
        attachment = self.env['ir.attachment'].create({
            'name': f'Наличие права на выполнение работ {registration_number}.docx',
            'datas': base64.b64encode(buffer.getvalue()),
            'res_model': self._name,
            'res_id': self[0].id,
            'type': 'binary',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }

class SroContactsDiscipline(models.Model):
    _name = 'sro.contacts.discipline'
    _description = "sro contacts discipline"

    is_discipline = fields.Boolean(string="Сведений нет")
    discipline_message = fields.Char(compute="_compute_discipline_message", store=False)

    disciplinary_basis = fields.Char(string="Основание открытия дисциплинарного производства")
    disciplinary_start_date = fields.Date(string="Дата открытия дисциплинарного производства")
    disciplinary_decision = fields.Text(string="Решение дисциплинарной комиссии")
    disciplinary_decision_date = fields.Date(string="Дата решения дисциплинарной комиссии")

    partner3_id = fields.Many2one('res.partner', invisible=True)

    @api.depends('is_discipline')
    def _compute_discipline_message(self):
        for record in self:
            record.discipline_message = "Сведений о дисциплинарных производствах нет." if record.is_discipline else ""

    @api.onchange('is_discipline')
    def _onchange_is_discipline(self):
        if self.is_discipline:
            self.update({
                'disciplinary_basis': False,
                'disciplinary_start_date': False,
                'disciplinary_decision': False,
                'disciplinary_decision_date': False,
            })

    def action_export_discipline_docx(self):
        partner = self.mapped('partner3_id')

        if not partner:
            raise UserError("Не удалось определить партнера!")
        records = self.search([('partner3_id', '=', partner.id)])

        if not records:
            raise UserError("Нет данных для экспорта!")
        doc = Document()

        section = doc.sections[0]
        section.left_margin = Cm(2)  # Уменьшаем левое поле (по умолчанию оно может быть 2 см или больше)
        section.right_margin = Cm(1)  # Правое поле оставляем стандартным или по желанию
        section.top_margin = Cm(2)  # Верхнее поле
        section.bottom_margin = Cm(2)  # Нижнее поле

        column_widths = [Cm(7), Cm(5), Cm(4), Cm(6), Cm(4)]

        table = doc.add_table(rows=1, cols=len(column_widths))

        def set_font(cell, bold=False):
            for para in cell.paragraphs:
                if not para.runs:
                    para.add_run()
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = bold

        registration_number = self.partner3_id.registration_number or "записи"
        headers = [
            "Заголовок",
            "Основание открытия дисциплинарного производства", 
            "Дата открытия дисциплинарного производства", 
            "Решение Дисциплинарной комиссии", 
            "Дата решения Дисциплинарной комиссии"
        ]

        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            set_font(cell, bold=True)

        for record in records:
            row_cells = table.add_row().cells
            row_cells[0].text = f"Дисциплинарное производство {registration_number}"
            row_cells[1].text = record.disciplinary_basis or "—"
            row_cells[2].text = record.disciplinary_start_date.strftime('%d.%m.%Y') if record.disciplinary_start_date else "—"
            row_cells[3].text = record.disciplinary_decision or "—"
            row_cells[4].text = record.disciplinary_decision_date.strftime('%d.%m.%Y') if record.disciplinary_decision_date else "—"

            for i, cell in enumerate(row_cells):
                set_font(cell)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        attachment = self.env['ir.attachment'].create({
            'name': f'Сведения о дисциплинарных взысканиях {registration_number}.docx',
            'datas': base64.b64encode(buffer.getvalue()),
            'res_model': self._name,
            'res_id': self[0].id,
            'type': 'binary',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }

class SroContactsInspection(models.Model):
    _name = 'sro.contacts.inspection'
    _description = "sro contacts inspection"
    _order = "inspection_number asc"
    _rec_name = 'inspection_name'

    inspections_conducted = fields.Boolean(string="Проверки не проводились")
    inspection_number = fields.Integer(string="№", index=True)
    inspection_short_name = fields.Char(string="Сокращенное наименование организации")
    inspection_name = fields.Char(string="Заголовок:")
    inspection_name_link = fields.Html(compute="_compute_inspection_name_link")
    inspection_member_number_link = fields.Html(compute="_compute_inspection_member_number_link", string="Регистрационный номер (ссылка)")
    inspection_member_number = fields.Selection(
        selection=lambda self: self._get_registration_numbers(),
        string="Регистрационный номер члена")
    inspection_act_date = fields.Date(string="Дата акта проверки")
    inspection_month_year = fields.Char(string="Месяц, год проверки")
    inspection_type = fields.Selection([
        ('scheduled', 'Плановая'),
        ('unscheduled', 'Внеплановая'),
    ], string="Вид проверки")
    inspection_form = fields.Selection([
        ('documentary', 'Документарная'),
        ('onsite', 'Выездная'),
    ], string="Форма проверки")
    inspection_law_violations = fields.Text(string="Нарушения требований законодательства РФ о градостроительной деятельности, о техническом регулировании, Стандартов НОСТРОЙ")
    inspection_internal_violations = fields.Text(string="Нарушения внутренних документов и стандартов А «СО «СЧ»")
    inspection_contract_violations = fields.Text(string="Нарушение исполнения обязательств по договорам строительного подряда, заключенным с использованием конкурентных способов заключения договоров")
    inspection_result = fields.Selection([
        ('yes', 'Нарушения имеются'),
        ('no', 'Нарушений не имеется'),
        ('draft', 'Проверки не проводились'),
    ], string="Результат проверки (итог)")
    inspection_disciplinary_measures = fields.Text(string="Применение мер дисциплинарного воздействия (итог)")
    inspection_measures_list = fields.Text(string="Перечень мер дисциплинарного воздействия")

    partner4_id = fields.Many2one('res.partner')

    def name_get(self):
        result = []
        for record in self:
            name = record.inspection_name or "Без названия"
            result.append((record.id, name))
        return result

    @api.depends("inspection_name")
    def _compute_inspection_name_link(self):
        for rec in self:
            if rec.id:  # Проверяем, что у записи есть ID
                url = f"/web#id={rec.id}&cids=1&menu_id=184&action=235&model=sro.contacts.inspection&view_type=form"
                rec.inspection_name_link = f'<a href="{url}" target="_blank">{rec.inspection_name}</a>'
            else:
                rec.inspection_name_link = ""

    @api.depends("inspection_member_number")
    def _compute_inspection_member_number_link(self):
        for rec in self:
            if rec.inspection_member_number:
                partner = self.env['res.partner'].search([('registration_number', '=', rec.inspection_member_number)], limit=1)
                if partner:
                    base_url = self.env['ir.config_parameter'].sudo().get_param('web.base.url')
                    url = f"{base_url}/web#id={partner.id}&model=res.partner&view_type=form"
                    rec.inspection_member_number_link = f'<a href="{url}" target="_blank">{rec.inspection_member_number}</a>'
                else:
                    rec.inspection_member_number_link = rec.inspection_member_number
            else:
                rec.inspection_member_number_link = ""

    @api.model
    def create(self, vals):
        """Присваиваем новой записи номер 1 и сдвигаем остальные"""
        records = self.search([], order="inspection_number asc")

        new_record = super(SroContactsInspection, self).create(vals)
        new_record.inspection_number = 1  # Новая запись всегда первая

        # Сдвигаем все остальные номера
        number = 2
        for record in records.sorted(key=lambda r: r.inspection_number):
            record.write({'inspection_number': number})
            number += 1

        return new_record

    def unlink(self):
        """Пересчитываем номера после удаления"""
        res = super(SroContactsInspection, self).unlink()

        # Перенумеровка оставшихся записей в правильном порядке
        records = self.search([], order="inspection_number asc")
        for index, record in enumerate(records.sorted(key=lambda r: r.inspection_number), start=1):
            record.write({'inspection_number': index})

        return res

    @api.model
    def _get_registration_numbers(self):
        partners = self.env['res.partner'].search([('registration_number', '!=', False)])
        return [(partner.registration_number, partner.registration_number) for partner in partners]

    @api.onchange('inspection_member_number')
    def _onchange_inspection_member_number(self):
        """Автоматически заполняет `inspection_short_name` по `inspection_member_number`"""
        if self.inspection_member_number:
            partner = self.env['res.partner'].search([('registration_number', '=', self.inspection_member_number)], limit=1)
            self.inspection_short_name = partner.short_name if partner else ''
        else:
            self.inspection_short_name = ''

    @api.onchange('partner4_id')
    def _onchange_partner_id(self):
        """При выборе партнера автоматически заполняет регистрационный номер, но поле остается редактируемым"""
        if self.partner4_id:
            self.inspection_member_number = self.partner4_id.registration_number
        else:
            self.inspection_member_number = False

    @api.onchange('inspections_conducted')
    def _onchange_inspections_conducted(self):
        """Очищает поля при снятии галочки."""
        if self.inspections_conducted:
            self.update({
                'inspection_member_number': False,
                'inspection_act_date': False,
                'inspection_month_year': False,
                'inspection_type': False,
                'inspection_form': False,
                'inspection_law_violations': False,
                'inspection_internal_violations': False,
                'inspection_contract_violations': False,
                'inspection_disciplinary_measures': False,
                'inspection_measures_list': False,
                'inspection_result': 'draft',
            })

    def action_export_inspection_docx(self):
        partner = self.mapped('partner4_id')

        if not partner:
            raise UserError("Не удалось определить партнера!")
        records = self.search([('partner4_id', '=', partner.id)])

        if not records:
            raise UserError("Нет данных для экспорта!")
        doc = Document()

        column_widths = [Cm(3), Cm(4), Cm(4), Cm(6), Cm(5)]

        table = doc.add_table(rows=1, cols=len(column_widths))

        def set_font(cell, bold=False):
            for para in cell.paragraphs:
                if not para.runs:
                    para.add_run()
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = bold

        headers = [
            "Заголовок", 
            "Дата решения", 
            "Месяц, год проверки", 
            "Перечень мер дисциплинарного воздействия", 
            "Результат проверки"
        ]

        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            set_font(cell, bold=True)

        for record in records:
            row_cells = table.add_row().cells
            row_cells[0].text = record.inspection_name or "—"
            row_cells[1].text = record.inspection_act_date.strftime('%d.%m.%Y') if record.inspection_act_date else "—"
            row_cells[2].text = record.inspection_month_year or "—"
            row_cells[3].text = record.inspection_measures_list or "—"
            row_cells[4].text = dict(record._fields['inspection_result'].selection).get(record.inspection_result, "—")

            for i, cell in enumerate(row_cells):
                set_font(cell)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        registration_number = self.partner4_id.registration_number or "записи"

        attachment = self.env['ir.attachment'].create({
            'name': f'Сведения о результатах проведенных проверок {registration_number}.docx',
            'datas': base64.b64encode(buffer.getvalue()),
            'res_model': self._name,
            'res_id': self[0].id,
            'type': 'binary',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }

class SroContactsContract(models.Model):
    _name = 'sro.contacts.contract'
    _description = "sro contacts contract"

    tender_date = fields.Datetime(string="Дата публикации", default=lambda self: fields.Datetime.now())
    tender_type = fields.Selection([
        ('offer', 'Предложение подряда'),
        ('search', 'Поиск подряда'),
        ('draft', 'Преложений нет'),
    ], string="Предложение или поиск подряда")
    tender_description = fields.Text(string="Описание подряда")
    tender_contact_info = fields.Text(string="Контактные данные")
    tender_member_number = fields.Selection(
        selection=lambda self: self._get_registration_numbers(),
        string="Регистрационный номер члена",
        help="Выберите регистрационный номер из списка или введите вручную.")
    tender_short_name = fields.Selection(
        selection=lambda self: self._get_short_name(),
        string="Сокращенное наименование организации")

    partner5_id = fields.Many2one('res.partner')

    @api.model
    def create(self, vals):
        if not vals.get('tender_date'):
            vals['tender_date'] = fields.Datetime.now()
        return super(SroContactsContract, self).create(vals)

    @api.model
    def _get_registration_numbers(self):
        """Получаем список регистрационных номеров из res.partner"""
        partners = self.env['res.partner'].search([('registration_number', '!=', False)])
        return [(partner.registration_number, partner.registration_number) for partner in partners]

    @api.model
    def _get_short_name(self):
        """Получаем список сокращенных наименований из res.partner"""
        partners = self.env['res.partner'].search([('short_name', '!=', False)])
        return [(partner.short_name, partner.short_name) for partner in partners]

    @api.onchange('partner5_id')
    def _onchange_partner_id(self):
        """При выборе партнера автоматически заполняет регистрационный номер, но поле остается редактируемым"""
        if self.partner5_id:
            self.tender_member_number = self.partner5_id.registration_number
        else:
            self.tender_member_number = False

    @api.onchange('tender_type')
    def _onchange_tender_type(self):
        if self.tender_type == 'draft':
            self.update({
                'tender_description': False,
                'tender_contact_info': False,
                'tender_member_number': False,
                'tender_short_name': False,
            })

    def action_export_contract_docx(self):
        partner = self.mapped('partner5_id')

        if not partner:
            raise UserError("Не удалось определить партнера!")
        records = self.search([('partner5_id', '=', partner.id)])

        if not records:
            raise UserError("Нет данных для экспорта!")
        doc = Document()

        column_widths = [Cm(4), Cm(6), Cm(5), Cm(4), Cm(5)]

        table = doc.add_table(rows=1, cols=len(column_widths))

        def set_font(cell, bold=False):
            for para in cell.paragraphs:
                if not para.runs:
                    para.add_run()
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = bold

        headers = [
            "Предложение или поиск подряда", 
            "Описание подряда", 
            "Контактные данные", 
            "Регистрационный номер члена", 
            "Сокращенное наименование организации"
        ]

        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            set_font(cell, bold=True)

        for record in records:
            row_cells = table.add_row().cells
            row_cells[0].text = dict(record._fields['tender_type'].selection).get(record.tender_type, "—")
            row_cells[1].text = record.tender_description or "—"
            row_cells[2].text = record.tender_contact_info or "—"
            row_cells[3].text = record.tender_member_number or "—"
            row_cells[4].text = record.tender_short_name or "—"

            for i, cell in enumerate(row_cells):
                set_font(cell)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        registration_number = self.partner5_id.registration_number or "записи"

        attachment = self.env['ir.attachment'].create({
            'name': f'Предложение подряда {registration_number}.docx',
            'datas': base64.b64encode(buffer.getvalue()),
            'res_model': self._name,
            'res_id': self[0].id,
            'type': 'binary',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }

class ConstructionRightSuspension(models.Model):
    _name = "sro.contacts.construction"
    _description = "sro contacts construction"

    number = fields.Char(string="№:")
    decision_date = fields.Date(string="Дата решения:")
    management_decision = fields.Selection([
        ('suspended', 'Приостановление действия права'),
        ('resumed', 'Возобновлено действия права'),
    ], string="Решение органов управления (право):")
    decision_basis = fields.Text(string="Основание решения:")

    partner6_id = fields.Many2one('res.partner')

    def action_export_construction_docx(self):
        partner = self.mapped('partner6_id')

        if not partner:
            raise UserError("Не удалось определить партнера!")
        records = self.search([('partner6_id', '=', partner.id)])

        if not records:
            raise UserError("Нет данных для экспорта!")
        doc = Document()

        section = doc.sections[0]
        section.left_margin = Cm(2)  # Уменьшаем левое поле (по умолчанию оно может быть 2 см или больше)
        section.right_margin = Cm(0.5)  # Правое поле оставляем стандартным или по желанию
        section.top_margin = Cm(2)  # Верхнее поле
        section.bottom_margin = Cm(2)  # Нижнее поле

        column_widths = [Cm(1), Cm(4), Cm(5), Cm(7), Cm(0.5)]

        table = doc.add_table(rows=1, cols=5)

        table.autofit = True
        for i, width in enumerate(column_widths):
            table.columns[i].width = width
            table.columns[0].width = Cm(1)

        def set_font(cell, bold=False):
            for para in cell.paragraphs:
                if not para.runs:
                    para.add_run()
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = bold

        headers = [
            "№", 
            "Дата решения",
            "Решение органов управления (право)",
            "Основание решения"
        ]

        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            set_font(cell, bold=True)

        for record in records:
            row_cells = table.add_row().cells
            row_cells[0].text = record.number or "—"
            row_cells[1].text = record.decision_date.strftime('%d.%m.%Y') if record.decision_date else "—"
            row_cells[2].text = dict(record._fields['management_decision'].selection).get(record.management_decision, "—")
            row_cells[3].text = record.decision_basis or "—"

            for i, cell in enumerate(row_cells):
                set_font(cell)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        registration_number = self.partner6_id.registration_number or "записи"

        attachment = self.env['ir.attachment'].create({
            'name': f'Сведения о приостановлении_возобновлении действия права {registration_number}.docx',
            'datas': base64.b64encode(buffer.getvalue()),
            'res_model': self._name,
            'res_id': self[0].id,
            'type': 'binary',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }

class InsurerInfo(models.Model):
    _name = "insurer.info"
    _description = "insurer info"

    name = fields.Char(string="Наименование организации:", store=True)
    license_number = fields.Char(string="№ Лицензии:", store=True)
    address = fields.Text(string="Адрес:", store=True)
    phone_insurer = fields.Char(string="Контактные телефоны:", store=True)
    website = fields.Char(string="Веб сайт:", store=True)
    email = fields.Char(string="Электронная почта:", store=True)

class ResPartner(models.Model):
    _inherit = 'res.partner'

    # Регистрационные данные
    registration_number = fields.Char(string="Регистрационный номер:", index=True, unique=True)
    short_name = fields.Char(string="Сокращенное наименование организации:")
    full_name = fields.Char(string="Полное наименование организации:")
    inn = fields.Char(string="ИНН:")
    ogrn = fields.Char(string="ОГРН:")
    registration_date = fields.Date(string="Дата гос. регистрации ЮЛ/ИП:")

    # Сведения о членстве в СРО
    sro_membership_compliance = fields.Selection([
        ('active', 'Соответствует'),
        ('suspended', 'Не соответствует'),
    ], string="Сведения о соответствии члена СРО условиям членства, предусмотренным законодательством РФ и (или) внутренними документами СРО:")
    sro_membership_status = fields.Selection([
        ('active', 'Является членом'),
        ('suspended', 'Исключен'),
    ], string="Статус членства")
    sro_registration_date = fields.Date(string="Дата регистрации в реестре СРО (внесения сведений в реестр):")
    sro_admission_basis = fields.Many2one('blog.post', string="Основание приема в СРО:") 
    sro_admission_basis_link = fields.Html(string="Основание приема в СРО:", 
        compute="_compute_sro_admission_basis_link", store=True)
    show_termination_fields = fields.Boolean(compute="_compute_show_termination_fields", store=False)
    termination_date = fields.Date(string="Дата прекращения членства:")
    termination_reason = fields.Many2one('blog.post', string="Основание прекращения членства:")
    termination_reason_link = fields.Html(string="Основание прекращения членства:", 
        compute="_compute_termination_reason_link", store=True)
    termination_info = fields.Text(string="Сведения о прекращении членства:")

    # Компенсационные фонды
    compensation_fund_vv_amount = fields.Float(string="Сумма взноса в Компенсационный Фонд возмещения вреда (КФ ВВ) (руб.):")
    vv_responsibility_level = fields.Char(string="Уровень ответственности ВВ:")
    contract_work_cost = fields.Char(string="Стоимость работ по одному договору:")
    compensation_fund_odo_amount = fields.Float(string="Сумма взноса в Компенсационный Фонд обеспечения договорных обязательств (КФ ОДО) (руб.):")
    odo_responsibility_level = fields.Char(string="Уровень ответственности ОДО:")
    max_obligation_amount = fields.Char(string="Предельный размер обязательств по договорам, заключаемым с использованием конкурентных способов заключения договоров:")

    # Руководство
    executive_authority = fields.Char(string="Единоличный исполнительный орган/руководитель коллегиального испольнительного органа:")

    # Контактные данные
    phone_sro = fields.Char(string="Контактные телефоны:")
    custom_website = fields.Char(string="Веб сайт:")
    zip = fields.Char(string="Адрес (Индекс):")
    country_id = fields.Many2one('res.country', string="Адрес (Страна):")
    state_id = fields.Many2one('res.country.state', string="Адрес (Субъект РФ):")
    hood = fields.Char(string="Адрес (Район):")
    city = fields.Char(string="Адрес (Населённый пункт):")
    street = fields.Char(string="Адрес (Улица):")
    street2 = fields.Char(string="Адреc (Дом):")
    corps = fields.Char(string="Адрес (Корпус / строение):")
    premises = fields.Char(string="Адрес (Помещение):")

    # Страхование
    insurer_info = fields.Many2one('insurer.info', string="Сведения о страховщике") 

    insurer_name = fields.Char(string="Наименование организации:")
    insurer_license_number = fields.Char(string="№ Лицензии:")
    insurer_address = fields.Text(string="Адрес:")
    insurer_phone = fields.Char(string="Контактные телефоны:")
    insurer_website = fields.Char(string="Веб сайт:")
    insurer_email = fields.Char(string="Электронная почта:")

    insurance_contract_number = fields.Char(string="Номер договора страхования:")
    insurance_contract_expiry = fields.Char(string="Срок действия договора страхования:")
    insurance_amount = fields.Float(string="Страховая сумма (руб.):")

    # Связь One2many с sro.contacts
    sro_contact2_ids = fields.One2many('sro.contacts.work', 'partner2_id', string="Наличие права на выполнение работ")
    sro_contact3_ids = fields.One2many('sro.contacts.discipline', 'partner3_id', string="Сведения о дисциплинарных производствах")
    sro_contact4_ids = fields.One2many('sro.contacts.inspection', 'partner4_id', string="Сведения о результатах проведенных проверок")
    sro_contact5_ids = fields.One2many('sro.contacts.contract', 'partner5_id', string="Предложения подрядов")
    sro_contact6_ids = fields.One2many('sro.contacts.construction', 'partner6_id', string="Сведения о приостановлении/возобновлении действия права выполнять строительсвто, реконструкцию, капиатльный ремонт объектов капитального строительства")

    @api.onchange('insurer_info')
    def _onchange_insurer_info(self):
        if self.insurer_info:
            self.insurer_name = self.insurer_info.name
            self.insurer_license_number = self.insurer_info.license_number
            self.insurer_address = self.insurer_info.address
            self.insurer_phone = self.insurer_info.phone_insurer
            self.insurer_website = self.insurer_info.website
            self.insurer_email = self.insurer_info.email
        else:
            self.insurer_name = False
            self.insurer_license_number = False
            self.insurer_address = False
            self.insurer_phone = False
            self.insurer_website = False
            self.insurer_email = False

    @api.depends("sro_admission_basis")
    def _compute_sro_admission_basis_link(self):
        for rec in self:
            if rec.sro_admission_basis:
                url = f"/blog/resheniia-pravleniia-3/{rec.sro_admission_basis.id}"
                rec.sro_admission_basis_link = f'<a href="{url}" target="_blank">{rec.sro_admission_basis.name}</a>'
            else:
                rec.sro_admission_basis_link = ""

    @api.depends("termination_reason")
    def _compute_termination_reason_link(self):
        for rec in self:
            if rec.termination_reason:
                url = f"/blog/resheniia-pravleniia-3/{rec.termination_reason.id}"
                rec.termination_reason_link = f'<a href="{url}" target="_blank">{rec.termination_reason.name}</a>'
            else:
                rec.termination_reason_link = ""

    @api.depends('sro_membership_status')
    def _compute_show_termination_fields(self):
        for record in self:
            record.show_termination_fields = record.sro_membership_status == 'suspended'

    def name_get(self):
        result = []
        for partner in self:
            name = partner.registration_number if partner.registration_number else partner.name or "Без номера"
            result.append((partner.id, name))
        return result

    @api.model
    def name_search(self, name, args=None, operator='ilike', limit=100):
        """Поиск по регистрационному номеру"""
        args = args or []
        if name:
            partners = self.search([('registration_number', operator, name)] + args, limit=limit)
        else:
            partners = self.search(args, limit=limit)
        return partners.name_get()

    def action_export_contact_docx(self):
        """Генерирует DOCX-файл с информацией о текущем контакте"""
        self.ensure_one()

        doc = Document()
        heading = doc.add_paragraph('Информация')
        run = heading.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        heading.alignment = 1  # Выравнивание по центру

        # Создаем одну таблицу с двумя колонками
        table = doc.add_table(rows=0, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(7)
        table.columns[1].width = Cm(9)

        def add_info(attribute, value):
            if not value:
                return  # Пропускаем пустые поля
            if isinstance(value, (date, datetime)):
                value = value.strftime('%d.%m.%Y')
            elif isinstance(value, float):
                value = f"{value:,.2f}".replace(",", " ")
            elif not value:
                value = "—"

            row_cells = table.add_row().cells
            row_cells[0].text = attribute
            row_cells[1].text = str(value)
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

        data = [
            ("Регистрационный номер:", self.registration_number),
            ("Сокращенное наименование организации:", self.short_name),
            ("Полное наименование организации:", self.full_name),
            ("ИНН:", self.inn),
            ("ОГРН:", self.ogrn),
            ("Дата гос. регистрации ЮЛ/ИП:", self.registration_date),
            ("Сведения о соответствии члена СРО условиям членства, предусмотренным законодательством РФ и (или) внутренними документами СРО:", dict(self._fields['sro_membership_compliance'].selection).get(self.sro_membership_compliance, '—')),
            ("Статус членства:", dict(self._fields['sro_membership_status'].selection).get(self.sro_membership_status, '—')),
            ("Дата регистрации в реестре СРО (внесения сведений в реестр):", self.sro_registration_date),
            ("Основание приема в СРО:", self.sro_admission_basis.name if self.sro_admission_basis else "—"),
        ]
        for attribute, value in data:
            add_info(attribute, value)

        # Добавляем блок "Сведения о прекращении членства", если нужно
        if self.show_termination_fields:
            termination_data = [
                ("Дата прекращения членства:", self.termination_date),
                ("Основание прекращения членства:", self.termination_reason.name),
                ("Сведения о прекращении членства:", self.termination_info),
            ]
            for attribute, value in termination_data:
                add_info(attribute, value)

        additional_data = [
            ("Сумма взноса в Компенсационный Фонд возмещения вреда (КФ ВВ) (руб.):", self.compensation_fund_vv_amount),
            ("Уровень ответственности ВВ:", self.vv_responsibility_level),
            ("Стоимость работ по одному договору:", self.contract_work_cost),
            ("Сумма взноса в Компенсационный Фонд обеспечения договорных обязательств (КФ ОДО) (руб.):", self.compensation_fund_odo_amount),
            ("Уровень ответственности ОДО:", self.odo_responsibility_level),
            ("Предельный размер обязательств по договорам, заключаемым с использованием конкурентных способов заключения договоров:", self.max_obligation_amount),
            ("Единоличный исполнительный орган/руководитель коллегиального исполнительного органа:", self.executive_authority),
            ("Контактные телефоны:", self.phone_sro),
            ("Веб сайт:", self.custom_website),
            ("Адрес (Индекс):", self.zip),
            ("Адрес (Страна):", self.country_id.name),
            ("Адрес (Субъект РФ):", self.state_id.name),
            ("Адрес (Район):", self.hood),
            ("Адрес (Населённый пункт):", self.city),
            ("Адрес (Улица):", self.street),
            ("Адрес (Дом):", self.street2),
            ("Адрес (Корпус/строение):", self.corps),
            ("Адрес (Помещение):", self.premises),
        ]

        for attribute, value in additional_data:
            add_info(attribute, value)

        # Добавляем строку "Сведения о страховщике" с данными в правой колонке
        insurer_info = (
            f"Наименование организации: {self.insurer_name}\n"
            f"№ Лицензии: {self.insurer_license_number}\n"
            f"Адрес: {self.insurer_address}\n"
            f"Контактные телефоны: {self.insurer_phone}\n"
            f"Веб сайт: {self.insurer_website}\n"
            f"Электронная почта: {self.insurer_email}\n"  
        )
        add_info("Сведения о страховщике:", insurer_info)

        # Отдельно добавляем страховые поля в обычном порядке
        insurance_data = [
            ("Номер договора страхования:", self.insurance_contract_number),
            ("Срок действия договора страхования:", self.insurance_contract_expiry),
            ("Страховая сумма (руб.):", f"{self.insurance_amount:,.2f}".replace(",", " ")),
        ]
        for attribute, value in insurance_data:
            add_info(attribute, value)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        attachment = self.env['ir.attachment'].create({
            'name': f'Данные реестра по {self.registration_number}.docx',
            'datas': base64.b64encode(buffer.getvalue()),
            'res_model': 'res.partner',
            'res_id': self.id,
            'type': 'binary',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }
